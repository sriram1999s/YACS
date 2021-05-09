from docx import Document
import pandas as pd
from datetime import datetime
import numpy as np
import matplotlib.pyplot as plt
import os
import json
from docx.enum.text import WD_BREAK
import os.path

def task_1(log_dict, algo_name, x , y, z):
    duration = [x1 - x2 for (x1, x2) in zip(log_dict['end_time'].values(), log_dict['arrival_time'].values())]
    mean_time = np.mean(duration)
    median_time = np.median(duration)
    x.append(algo_name)
    y.append(mean_time)
    z.append(median_time)

    str1 = "\t1. Mean Completion Time : " + '{:.4f}'.format(mean_time) + " seconds\n"
    str2 = "\t2. Median Completion Time : " + '{:.4f}'.format(median_time) + " seconds\n"

    p_object = document.add_paragraph("")
    p_object.add_run(str1)
    p_object.add_run(str2)


def mean_median_dump(log_dict, algo_name):
    duration = [x1 - x2 for (x1, x2) in zip(log_dict['end_time'].values(), log_dict['arrival_time'].values())]
    mean_time = np.mean(duration)
    median_time = np.median(duration)

    str1 = "\t1. Mean Completion Time : " + '{:.4f}'.format(mean_time) + " seconds\n"
    str2 = "\t2. Median Completion Time : " + '{:.4f}'.format(median_time) + " seconds\n"

    str3 = 'Task Analysis\n'
    image_name = 'img/'+ algo_name + '_Task_Graph.png'

    p_object = document.add_paragraph("")
    # p_object.add_run(str3).bold = True
    p_object.add_run(str1)
    p_object.add_run(str2)

    y = [mean_time, median_time]
    x = ['mean completion time', 'median completion time']

    bar_graph = plt.bar(x, y,)
    bar_graph[0].set_color('m')
    bar_graph[1].set_color('y')

    plt.ylabel("Completion times")
    plt.xlabel("Measure of Central Tendency")
    plt.title("Completion times vs Measures of Central Tendency")
    plt.savefig(image_name)
    plt.show(block=False)
    plt.close()
    document.add_picture(image_name)
    document.add_page_break()

    str3 = 'Plot'
    p_object = document.add_paragraph("")
    p_object.add_run(str3).bold = True


def plotter(log_dict, algo_name):
    imgname = algo_name # Getting the name of the algorithm
    imgname += "_plot_image.png" # obtaining the appropriate name to save the plot later
    imgname = "img/" + imgname
    # doc_name = algo_name + "_Logs_Analysis.docx"
    info = [] # list of tuples [('arrival_time', 'worker_id')] in order to sort in ascending order of 'arrival time'
    unique_worker_set = set()
    check_if_done = dict()
    end_time_list = []

    for i, j in zip(log_dict['arrival_time'].values(), log_dict['worker_id'].values()):
        info.append((i, j)) # populating the list of tuples
        unique_worker_set.add(j)

    for i, j in zip(log_dict['end_time'].values(), log_dict['worker_id'].values()):
        check_if_done[i] = j
        end_time_list.append(i)

    graph_dict = dict()     # dictionary to store the number of tasks assigned to each worker every time a task arrives for scheduling
    graph_dict['time'] = list()
    graph_dict['date'] = list()
    for i in unique_worker_set:
        graph_dict[i] = list()

    info.sort(key = lambda x : x[0]) # sorting the list of tuples
    for j, i in info:
        worker_list = list(unique_worker_set) # set of all the workers
        arrival_time = str(datetime.fromtimestamp(j))
        date, time = arrival_time.split(' ')
        graph_dict['time'].append(time) # keying in the arrival time of each task
        graph_dict['date'].append(date) # keying in the arrival time of each task



        if graph_dict[i] == []:
            graph_dict[i] = [1] # when the first task of the job comes in, we initialize the task count of the worker that is alloted this task
        else:
            top = len(graph_dict[i]) - 1
            graph_dict[i].append(graph_dict[i][top] + 1) # updating the task count for the worker that was alloted this task

        worker_list.remove(i) # remove the worker that was alloted this task from the list of all workers
        for k in worker_list:
            top = len(graph_dict[k]) - 1
            if top == -1:
                graph_dict[k].append(0)
            else:
                graph_dict[k].append(graph_dict[k][top]) #keeping the new task count of the other workers same as the previously updated task count

        for end_time in end_time_list:
            if end_time <= j:
                worker_id = check_if_done[end_time]
                top = len(graph_dict[worker_id]) - 1
                graph_dict[worker_id][top] = graph_dict[worker_id][top] - 1
                end_time_list.remove(end_time)

    # plotting the graph
    x_axis = list(range(1, len(info) + 1))
    legend_text_list = []
    for i in graph_dict.keys():
        if i != 'time' and i != 'date':
            plt.plot(x_axis, graph_dict[i])
            legend_text_list.append('Worker ' + str(i))

    plt.xticks(rotation='vertical')
    plt.xlabel('Arrival time of a task')
    plt.ylabel('Number of tasks scheduled for each worker')
    # plt.legend(['worker 0', 'worker 1', 'worker 2'])
    plt.legend(legend_text_list)
    plt.title('Number of tasks scheduled on each machine against time')
    plt.savefig(imgname)
    plt.show(block=False)
    plt.close()

    para_object = document.add_paragraph('DATE\t\t\tARRIVAL TIME\t\t\tX-AXIS EQUIVALENTS\n')
    for i,j,k in zip(graph_dict['date'], graph_dict['time'], x_axis):
        para_object.add_run(str(i))
        para_object.add_run("\t\t")
        para_object.add_run(str(j))
        para_object.add_run("\t\t\t\t")
        para_object.add_run(str(k))
        para_object.add_run("\n")

    document.add_picture(imgname)
    document.add_page_break()

def print_bar_graph(x, y, y_name):

    document.add_heading(y_name + " Comparison\n\n", level=1)

    str3 = 'Job Analysis\n'
    image_name = 'img/'+ y_name + '_Job_Graph.png'

    colors = ['r' , 'g', 'b']
    bar_graph = plt.bar(x, y )
    for i in range(len(x)):
        bar_graph[i].set_color(colors[i])
    plt.ylabel(y_name)
    plt.xlabel("Algorithms")
    plt.title(y_name + " of Completion vs Algorithms")
    plt.savefig(image_name)
    plt.show(block=False)
    plt.close()
    document.add_picture(image_name)
    document.add_page_break()


if __name__=="__main__":

    current_path = os.getcwd()
    folder_path = os.path.join(current_path,'img')
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    document = Document()
    doc_name = "Log_analysis.docx"
    doc_heading = "Analysis Of YACS Logs"
    document.add_heading(doc_heading,0)
    document.add_heading("\nAnalysis of Jobs\n\n", level=1.5)
    # Analysing logs/job_logs.csv
    df1 = pd.read_csv("logs/job_log.csv")
    df_algos1={}
    df_algos1['RANDOM'] = df1[df1.algo == "RANDOM"]
    df_algos1['RR'] = df1[df1.algo == "RR"]
    df_algos1['LL'] = df1[df1.algo == "LL"]
    x=[]
    y=[]
    z=[]
    for key in df_algos1:

        if not df_algos1[key].empty:
            # Converting the dataframe 'df' into a dictonary 'job_log_dict'
            job_log_dict = df_algos1[key].to_dict()
            sub_heading = key +' Scheduling algorithm: '
            document.add_heading(sub_heading , level=2)
            task_1(job_log_dict, key, x, y ,z)
    print_bar_graph(x,y,"Mean")
    print_bar_graph(x,z,"Median")


    # Analysing logs/task_logs.csv
    df2 = pd.read_csv("logs/task_log.csv")
    df_algos2={}
    df_algos2['RANDOM'] = df2[df2.algo == "RANDOM"]
    df_algos2['RR'] = df2[df2.algo == "RR"]
    df_algos2['LL'] = df2[df2.algo == "LL"]
    document.add_heading("Analysis of Tasks", level=1)


    for key in df_algos2:
        if not df_algos2[key].empty:
            # Converting the dataframe 'df' into a dictonary 'task_log_dict'
            task_log_dict = df_algos2[key].to_dict()
            sub_heading = key +' Scheduling algorithm: '
            document.add_heading(sub_heading , level=2)
            mean_median_dump(task_log_dict, key)
            plotter(task_log_dict,key)

    document.save(doc_name)
